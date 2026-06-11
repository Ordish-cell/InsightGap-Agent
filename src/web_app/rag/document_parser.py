import csv
import json
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup

ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".xlsx", ".csv", ".json", ".html", ".htm"}


def parse_document(path: str | Path, original_filename: str | None = None, mime_type: str | None = None) -> dict[str, Any]:
    file_path = Path(path).resolve()
    if not file_path.exists() or not file_path.is_file():
        raise FileNotFoundError("Document file does not exist")
    ext = file_path.suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {ext}")

    fallback_reason = ""
    try:
        parsed = _parse_markitdown(file_path)
        if parsed["text"].strip():
            parsed["metadata"].update(_base_metadata(file_path, "markitdown", original_filename, mime_type))
            if ext == ".pdf":
                try:
                    _, pdf_metadata = _parse_pdf(file_path, dict(parsed["metadata"]))
                    parsed["metadata"].update({
                        "page_count": pdf_metadata.get("page_count", 0),
                        "pages": pdf_metadata.get("pages", []),
                        "page_parser": pdf_metadata.get("parser", ""),
                    })
                except Exception:
                    pass
            parsed["metadata"].update({"used_fallback": False, "fallback_reason": ""})
            return parsed
    except Exception as exc:
        fallback_reason = f"{type(exc).__name__}: {exc}"

    parser = "text"
    metadata = _base_metadata(file_path, parser, original_filename, mime_type)
    if ext in {".txt", ".md"}:
        text = _read_text(file_path)
        markdown = text
    elif ext == ".pdf":
        text, metadata = _parse_pdf(file_path, metadata)
        markdown = text
    elif ext == ".docx":
        text = _parse_docx(file_path)
        markdown = text
    elif ext == ".xlsx":
        text, metadata = _parse_xlsx(file_path, metadata)
        markdown = text
    elif ext == ".csv":
        text, metadata = _parse_csv(file_path, metadata)
        markdown = text
    elif ext == ".json":
        text = json.dumps(json.loads(_read_text(file_path)), ensure_ascii=False, indent=2)
        markdown = f"```json\n{text}\n```"
    elif ext in {".html", ".htm"}:
        soup = BeautifulSoup(_read_text(file_path), "html.parser")
        text = soup.get_text("\n")
        markdown = text
    else:
        raise ValueError(f"Unsupported document type: {ext}")
    metadata.setdefault("parser", parser)
    metadata["used_fallback"] = bool(fallback_reason)
    metadata["fallback_reason"] = fallback_reason
    return {"text": text, "markdown": markdown, "metadata": metadata}


def _parse_markitdown(file_path: Path) -> dict[str, Any]:
    from markitdown import MarkItDown

    result = MarkItDown().convert(str(file_path))
    text = getattr(result, "text_content", "") or ""
    return {"text": text, "markdown": text, "metadata": {"parser": "markitdown", "page_count": 0, "sheet_names": []}}


def _base_metadata(file_path: Path, parser: str, original_filename: str | None, mime_type: str | None) -> dict[str, Any]:
    return {"parser": parser, "page_count": 0, "sheet_names": [], "source_filename": original_filename or file_path.name, "mime_type": mime_type or ""}


def _read_text(file_path: Path) -> str:
    for encoding in ("utf-8", "gbk"):
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return file_path.read_text(encoding="utf-8", errors="ignore")


def _parse_pdf(file_path: Path, metadata: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    import fitz

    doc = fitz.open(str(file_path))
    metadata["parser"] = "pymupdf"
    metadata["page_count"] = doc.page_count
    pages = []
    for index, page in enumerate(doc, 1):
        page_text = page.get_text("text")
        pages.append({"page_number": index, "text": page_text})
    metadata["pages"] = pages
    text = "\n\n".join(page["text"] for page in pages)
    doc.close()
    return text, metadata


def _parse_docx(file_path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "DOCX parsing requires python-docx. Install it with: pip install python-docx"
        ) from None
    except Exception as exc:
        raise RuntimeError(
            f"Failed to initialize python-docx parser: {exc}"
        ) from exc
    try:
        doc = Document(str(file_path))
        lines: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = getattr(paragraph.style, "name", "") or ""
            if style_name.lower().startswith("heading"):
                lines.append(f"# {text}")
            else:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(values):
                    lines.append("\t".join(values))
        return "\n".join(lines)
    except Exception as exc:
        raise RuntimeError(f"Failed to parse DOCX file: {exc}") from exc


def _parse_xlsx(file_path: Path, metadata: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError(
            "XLSX parsing requires openpyxl. Install it with: pip install openpyxl"
        ) from None
    except Exception as exc:
        raise RuntimeError(
            f"Failed to initialize openpyxl parser: {exc}"
        ) from exc
    workbook = openpyxl.load_workbook(str(file_path), data_only=True, read_only=True)
    metadata["parser"] = "openpyxl"
    metadata["sheet_names"] = workbook.sheetnames
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"# {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                lines.append("\t".join(values))
    workbook.close()
    return "\n".join(lines), metadata


def _parse_csv(file_path: Path, metadata: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    with file_path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        rows = list(csv.reader(handle))
    metadata["parser"] = "csv"
    metadata["row_count"] = max(0, len(rows) - 1)
    metadata["column_count"] = len(rows[0]) if rows else 0
    return "\n".join("\t".join(row) for row in rows), metadata
